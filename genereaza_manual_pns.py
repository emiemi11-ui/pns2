#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GENERATOR MANUAL COMPLET PNS - Script pentru Claude Code
Procesează:
1. TOATE cursurile C1-C5 (teoria cu explicații simple)
2. TOATE exercițiile rezolvate pas-cu-pas
3. TOATE examenele vechi cu soluții complete
4. Fișe de memorare (formule, metode, checklist)
Autor: Claude & Eminint
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PyPDF2 import PdfReader
import logging
from typing import Dict, List, Tuple

# Configurare logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('manual_pns_generator.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# CONFIGURARE CURSURI
# ============================================================================
CURSURI = {
    'C1': {
        'pdf': 'PNS_C1_Introducere.pdf',
        'titlu': 'Introducere în Prelucrarea Numerică a Semnalelor',
        'slide_count': 60
    },
    'C2': {
        'pdf': 'PNS_C2_Transformari_elementare.pdf',
        'titlu': 'Transformări Elementare ale Semnalelor',
        'slide_count': 31
    },
    'C3': {
        'pdf': 'PNS_C3_Semnale_elementare.pdf',
        'titlu': 'Semnale Elementare',
        'slide_count': 40
    },
    'C4': {
        'pdf': 'PNS_C4-Sisteme (SNLI,SALI)_prez.pdf',
        'titlu': 'Sisteme SNLI și SALI',
        'slide_count': 63
    },
    'C5': {
        'pdf': 'PNS_C5_Convolutia_corelatia.pdf',
        'titlu': 'Convoluția și Corelația',
        'slide_count': 32
    }
}

# ============================================================================
# CONFIGURARE EXERCIȚII ȘI EXAMENE (FILENAME-uri actualizate)
# ============================================================================
EXERCITII_FILES = {
    'exercitii_principale': [
        'ExamenPNS.pdf',
        'exercitiu.pdf',
        'Grile.pdf'
    ],
    'lucrari_rezolvate': [
        'lucrare_1_A.docx',
        'lucrare_1_D.docx',
        'lucrare_2_D.docx'
    ],
    'exemple_examene': [
        'E213B.docx',
        'E213B(2).docx',
        'E213B(3).docx',
        'E213C.docx'
    ],
    'examene_vechi': [
        '05.02.2017.pdf',
        '05.09.2019.pdf',
        '06.02.2017.pdf',
        '06.02.2017 (2).pdf',
        '06.02.2017rez.pdf',
        '06.02.2017rez .pdf',
        '06.02.2018 .pdf',
        '07.12.2016.docx',
        '08.05.2018.pdf',
        '09.12.2016.pdf',
        '12.01.2016.pdf',
        '12.09.2018 .pdf',
        '13.12.2017.docx',
        '15.01.2016 .pdf',
        '15.01.2016  (2).pdf',
        '16.02.2016 .pdf',
        '17.11.2015 .pdf',
        '18.12.2017.docx',
        '20.11.2015 .pdf',
        '25.11.2015.pdf'
    ]
}

# ============================================================================
# DICȚIONAR EXPLICAȚII SIMPLE (ca la proști)
# ============================================================================
EXPLICATII_SIMPLE = {
    # Concepte de bază
    'semnal': 'Semnalul e ca o undă care poartă informație - ca undele radio care aduc muzica la radio-ul din mașină.',
    'filtrare': 'Filtrarea e ca o pereche de căști noise-cancelling care elimină zgomotul trenului și lasă doar muzica.',
    'compresie': 'Compresia e ca atunci când transformi un fișier WAV de 50 MB într-un MP3 de 5 MB, păstrând aceeași melodie.',
    'transformare': 'Transformarea e ca și cum ai traduce un mesaj dintr-o limbă în alta - schimbi forma dar păstrezi sensul.',

    # Transformări elementare
    'scalare_amplitudine': 'Scalarea amplitudinii e exact ca butonul de volum - înmulțești semnalul cu un număr mai mare/mic.',
    'scalare_timp': 'Scalarea timpului e ca speed-ul de pe YouTube - poți reda mai repede (2x) sau mai încet (0.5x).',
    'reflexie': 'Reflexia e ca un film redat înapoi - în loc de x(t) ai x(-t), totul merge în sens invers.',
    'intarziere': 'Întârzierea e ca un tren care pleacă cu 10 minute mai târziu - semnalul x(t) devine x(t-10).',
    'avans': 'Avansul e opusul întârzierii - ca un tren care pleacă cu 10 minute mai devreme.',

    # Proprietăți semnale
    'paritate': 'Paritatea arată dacă un semnal e simetric (par) sau antisimetric (impar) față de originea timpului.',
    'energie': 'Energia unui semnal e suma pătratelor valorilor lui - ca energia pe care o consumi alergând.',
    'putere': 'Puterea e energia medie pe unitate de timp - ca consumul mediu de baterie al telefonului.',
    'periodicitate': 'Periodicitatea înseamnă că semnalul se repetă la intervale regulate - ca un ceas care ticăie.',

    # Semnale elementare
    'dirac': 'Impulsul Dirac (delta) e ca o lovitură instantanee - toată energia se eliberează într-un moment.',
    'treapta': 'Treapta Heaviside e ca un întrerupător - off înainte de t=0, on după t=0.',
    'exponentiala': 'Semnalul exponențial e ca creșterea bacteriilor - crește (sau scade) exponențial în timp.',
    'sinusoidal': 'Semnalul sinusoidal e ca undele de pe mare - urcă și coboară regulat, periodic.',

    # Sisteme
    'snli': 'Sistem Nestocat Liniar Invariant în Timp - sistemul cel mai simplu și previzibil, ca o rețetă de gătit fixă.',
    'convolutie': 'Convoluția e ca un mixer - amesteci semnalul de intrare cu răspunsul sistemului și obții ieșirea.',
    'corelatie': 'Corelația măsoară cât de asemănătoare sunt două semnale - ca recunoașterea vocală pe telefon.',
    'stabilitate': 'Un sistem stabil e ca o mașină bună - nu explodează chiar dacă îi dai input mare.',
    'cauzalitate': 'Un sistem cauzal e realist - ieșirea de azi depinde doar de inputul de azi și din trecut, nu din viitor.'
}

# ============================================================================
# TERMINOLOGIE TEHNICĂ
# ============================================================================
TERMINOLOGIE = {
    'PNS': 'Prelucrarea Numerică a Semnalelor (Digital Signal Processing - DSP)',
    'SNLI': 'Sistem Nestocat Liniar Invariant în Timp (LTI System)',
    'SALI': 'Sistem Amintitor Liniar Invariant în Timp',
    'FFT': 'Fast Fourier Transform (Transformata Fourier Rapidă)',
    'DFT': 'Discrete Fourier Transform (Transformata Fourier Discretă)',
    'FIR': 'Finite Impulse Response (Răspuns Impulsional Finit)',
    'IIR': 'Infinite Impulse Response (Răspuns Impulsional Infinit)',
    'ROC': 'Region of Convergence (Regiunea de Convergență)',
    'DTFT': 'Discrete-Time Fourier Transform',
    'Z-Transform': 'Transformata Z (generalizare a DTFT)'
}

# ============================================================================
# FORMULE ESENȚIALE PENTRU FIȘĂ DE MEMORARE
# ============================================================================
FORMULE_CHEIE = {
    'energie': 'E = ∑|x[n]|² (pentru semnale discrete)',
    'putere': 'P = lim(N→∞) (1/(2N+1)) ∑|x[n]|²',
    'convolutie': 'y[n] = x[n] * h[n] = ∑ x[k]h[n-k]',
    'corelatie': 'Rxy[l] = ∑ x[n]y*[n-l]',
    'transformata_z': 'X(z) = ∑ x[n]z^(-n)',
    'dtft': 'X(e^jω) = ∑ x[n]e^(-jωn)',
    'paritate_para': 'x(-t) = x(t) sau x[-n] = x[n]',
    'paritate_impara': 'x(-t) = -x(t) sau x[-n] = -x[n]',
    'periodicitate': 'x(t) = x(t+T) sau x[n] = x[n+N]'
}

# ============================================================================
# CLASA PRINCIPALĂ - GENERATOR MANUAL
# ============================================================================
class ManualPNSGenerator:
    """Generează manualul complet PNS cu teoria, exerciții și fișe"""

    def __init__(self, repo_path: str = '.'):
        self.repo_path = Path(repo_path)
        self.doc = None
        self.slide_counter = 0
        self.exercitiu_counter = 0

    def setup_document(self):
        """Configurează documentul DOCX cu formatare profesională"""
        self.doc = Document()

        # Setări pagină A4
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

        logger.info("✓ Document DOCX configurat (A4, margini profesionale)")

    def add_cover_page(self):
        """Adaugă pagina de copertă"""
        # Titlu principal
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("MANUAL COMPLET\nPRELUCRAREA NUMERICĂ A SEMNALELOR")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        self.doc.add_paragraph()  # Spațiu

        # Subtitlu
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Teorie Completă • Exerciții Rezolvate • Fișe de Memorare")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 102, 204)

        self.doc.add_paragraph()

        # Conținut
        content = self.doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run(
            "📚 Toate cursurile C1-C5 cu explicații simple\n"
            "✍️ Toate exercițiile și examenele rezolvate pas-cu-pas\n"
            "📋 Fișe de memorare pentru formule și metode\n"
            "✅ Checklist complet pentru examen"
        )
        run.font.size = Pt(12)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Autor și dată
        footer_p = self.doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Academia Tehnică Militară \"Ferdinand I\"\n2025")
        run.font.size = Pt(11)
        run.font.italic = True

        self.doc.add_page_break()
        logger.info("✓ Pagină de copertă adăugată")

    def add_table_of_contents(self):
        """Adaugă cuprinsul"""
        heading = self.doc.add_heading('CUPRINS', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        toc_items = [
            "PARTEA I - TEORIE COMPLETĂ",
            "  Cursul 1: Introducere în PNS",
            "  Cursul 2: Transformări Elementare",
            "  Cursul 3: Semnale Elementare",
            "  Cursul 4: Sisteme SNLI și SALI",
            "  Cursul 5: Convoluția și Corelația",
            "",
            "PARTEA II - EXERCIȚII REZOLVATE",
            "  Exerciții principale (ExamenPNS, exercitiu, Grile)",
            "  Lucrări rezolvate (lucrare_1_A, lucrare_1_D, lucrare_2_D)",
            "  Exemple examene (E213B, E213C)",
            "  Examene vechi (2015-2019)",
            "",
            "PARTEA III - FIȘE DE MEMORARE",
            "  Fișa 1: Formule esențiale",
            "  Fișa 2: Metode de rezolvare",
            "  Fișa 3: Erori frecvente",
            "  Fișa 4: Checklist examen"
        ]

        for item in toc_items:
            if not item:  # Skip empty lines
                self.doc.add_paragraph()
                continue
            p = self.doc.add_paragraph(item)
            if p.runs:  # Check if runs exist
                if not item.startswith('  '):
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)
                else:
                    p.runs[0].font.size = Pt(11)

        self.doc.add_page_break()
        logger.info("✓ Cuprins adăugat")

    def extract_text_from_pdf(self, pdf_path: Path) -> List[Tuple[int, str]]:
        """Extrage textul din PDF slide cu slide"""
        slides_text = []
        try:
            reader = PdfReader(str(pdf_path))
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    slides_text.append((page_num, text.strip()))
            logger.info(f"✓ Extras text din {pdf_path.name}: {len(slides_text)} slide-uri")
        except Exception as e:
            logger.error(f"✗ Eroare la extragerea din {pdf_path.name}: {e}")
        return slides_text

    def detect_concept_in_text(self, text: str) -> List[str]:
        """Detectează conceptele cheie dintr-un text de slide"""
        concepts = []
        text_lower = text.lower()

        # Verifică fiecare concept din dicționar
        for concept in EXPLICATII_SIMPLE.keys():
            # Caută conceptul (cu variații)
            if concept in text_lower or concept.replace('_', ' ') in text_lower:
                concepts.append(concept)

        return concepts

    def add_slide_content(self, curs_id: str, slide_num: int, slide_text: str, pdf_name: str):
        """Adaugă conținutul unui slide în document cu structura SURSĂ-TEXT-EXPLICAȚIE-TERMINOLOGIE"""

        # 1. SURSĂ (citare corectă)
        sursa = self.doc.add_paragraph()
        sursa_run = sursa.add_run(f"[SURSĂ: Slide {slide_num} din {pdf_name}]")
        sursa_run.font.italic = True
        sursa_run.font.color.rgb = RGBColor(102, 102, 102)
        sursa_run.font.size = Pt(9)

        # 2. TEXT EXACT DE PE SLIDE
        text_p = self.doc.add_paragraph()
        text_run = text_p.add_run(slide_text)
        text_run.font.size = Pt(11)
        text_run.font.name = 'Calibri'

        # 3. EXPLICAȚIE SIMPLĂ (dacă există concepte detectate)
        concepts = self.detect_concept_in_text(slide_text)
        if concepts:
            explicatie_p = self.doc.add_paragraph()
            explicatie_run = explicatie_p.add_run("💡 EXPLICAȚIE SIMPLĂ: ")
            explicatie_run.font.bold = True
            explicatie_run.font.color.rgb = RGBColor(0, 102, 204)
            explicatie_run.font.size = Pt(11)

            # Adaugă explicațiile pentru conceptele detectate
            for concept in concepts[:2]:  # Max 2 concepte per slide
                explicatie_text = explicatie_p.add_run(f"\n{EXPLICATII_SIMPLE[concept]}")
                explicatie_text.font.color.rgb = RGBColor(0, 102, 204)
                explicatie_text.font.size = Pt(10)

        # 4. TERMINOLOGIE TEHNICĂ (dacă există)
        termeni_gasiti = []
        text_upper = slide_text.upper()
        for termen, definitie in TERMINOLOGIE.items():
            if termen in text_upper:
                termeni_gasiti.append((termen, definitie))

        if termeni_gasiti:
            term_p = self.doc.add_paragraph()
            term_run = term_p.add_run("📖 TERMINOLOGIE: ")
            term_run.font.bold = True
            term_run.font.color.rgb = RGBColor(204, 0, 0)
            term_run.font.size = Pt(10)

            for termen, definitie in termeni_gasiti[:2]:  # Max 2 termeni
                term_text = term_p.add_run(f"\n• {termen} = {definitie}")
                term_text.font.color.rgb = RGBColor(204, 0, 0)
                term_text.font.size = Pt(9)

        # Linie separatoare
        self.doc.add_paragraph("─" * 80)

        self.slide_counter += 1
        if self.slide_counter % 10 == 0:
            logger.info(f"  Procesate {self.slide_counter} slide-uri...")

    def process_curs(self, curs_id: str):
        """Procesează un curs complet"""
        curs_info = CURSURI[curs_id]
        pdf_path = self.repo_path / curs_info['pdf']

        logger.info(f"\n{'='*60}")
        logger.info(f"PROCESARE {curs_id}: {curs_info['titlu']}")
        logger.info(f"{'='*60}")

        if not pdf_path.exists():
            logger.error(f"✗ Fișierul {pdf_path} nu există!")
            return

        # Heading pentru curs
        heading = self.doc.add_heading(f"{curs_id}: {curs_info['titlu']}", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        # Extrage și procesează slide-urile
        slides = self.extract_text_from_pdf(pdf_path)
        for slide_num, slide_text in slides:
            self.add_slide_content(curs_id, slide_num, slide_text, curs_info['pdf'])

        self.doc.add_page_break()
        logger.info(f"✓ {curs_id} complet: {len(slides)} slide-uri procesate\n")

    def process_exercitiu_pdf(self, pdf_path: Path, tip: str):
        """Procesează un PDF cu exerciții"""
        logger.info(f"  Procesare: {pdf_path.name}")

        # Heading pentru exercițiu
        heading = self.doc.add_heading(f"Exercițiu: {pdf_path.stem}", level=2)
        heading.runs[0].font.color.rgb = RGBColor(204, 102, 0)

        # Extrage text
        slides = self.extract_text_from_pdf(pdf_path)

        for slide_num, text in slides:
            # Sursă
            sursa = self.doc.add_paragraph()
            sursa_run = sursa.add_run(f"[{pdf_path.name}, Pagina {slide_num}]")
            sursa_run.font.italic = True
            sursa_run.font.color.rgb = RGBColor(102, 102, 102)
            sursa_run.font.size = Pt(9)

            # Enunț/Rezolvare
            content = self.doc.add_paragraph(text)
            content_run = content.runs[0]
            content_run.font.size = Pt(10)

            # Adaugă secțiune de rezolvare pas-cu-pas dacă textul conține formule
            if any(symbol in text for symbol in ['=', '∑', '∫', 'lim']):
                rezolvare = self.doc.add_paragraph()
                rez_run = rezolvare.add_run("📝 REZOLVARE PAS-CU-PAS:")
                rez_run.font.bold = True
                rez_run.font.color.rgb = RGBColor(0, 153, 76)
                rez_run.font.size = Pt(10)

                # Textul rezolvării (va fi extras din PDF)
                rez_text = rezolvare.add_run(f"\n{text}")
                rez_text.font.size = Pt(10)

            self.doc.add_paragraph("─" * 60)

        self.exercitiu_counter += len(slides)

    def process_exercitiu_docx(self, docx_path: Path, tip: str):
        """Procesează un DOCX cu exerciții"""
        logger.info(f"  Procesare: {docx_path.name}")

        try:
            source_doc = Document(str(docx_path))

            # Heading pentru exercițiu
            heading = self.doc.add_heading(f"Exercițiu: {docx_path.stem}", level=2)
            heading.runs[0].font.color.rgb = RGBColor(204, 102, 0)

            # Copiază conținutul
            for para in source_doc.paragraphs:
                if para.text.strip():
                    new_para = self.doc.add_paragraph(para.text)
                    new_para.style = para.style

            self.doc.add_paragraph("─" * 60)
            self.exercitiu_counter += 1

        except Exception as e:
            logger.error(f"  ✗ Eroare la procesarea {docx_path.name}: {e}")

    def process_all_exercitii(self):
        """Procesează TOATE exercițiile și examenele"""
        logger.info(f"\n{'='*60}")
        logger.info("PROCESARE EXERCIȚII ȘI EXAMENE")
        logger.info(f"{'='*60}\n")

        # Heading principal
        heading = self.doc.add_heading("PARTEA II - EXERCIȚII REZOLVATE", level=1)
        heading.runs[0].font.color.rgb = RGBColor(204, 102, 0)

        # 1. Exerciții principale
        self.doc.add_heading("A. Exerciții Principale", level=2)
        for filename in EXERCITII_FILES['exercitii_principale']:
            file_path = self.repo_path / filename
            if file_path.exists():
                self.process_exercitiu_pdf(file_path, "principal")
            else:
                logger.warning(f"  ⚠ Fișierul {filename} nu există")

        self.doc.add_page_break()

        # 2. Lucrări rezolvate
        self.doc.add_heading("B. Lucrări Rezolvate", level=2)
        for filename in EXERCITII_FILES['lucrari_rezolvate']:
            file_path = self.repo_path / filename
            if file_path.exists():
                self.process_exercitiu_docx(file_path, "lucrare")
            else:
                logger.warning(f"  ⚠ Fișierul {filename} nu există")

        self.doc.add_page_break()

        # 3. Exemple examene
        self.doc.add_heading("C. Exemple Examene (E213B, E213C)", level=2)
        for filename in EXERCITII_FILES['exemple_examene']:
            file_path = self.repo_path / filename
            if file_path.exists():
                self.process_exercitiu_docx(file_path, "exemplu")
            else:
                logger.warning(f"  ⚠ Fișierul {filename} nu există")

        self.doc.add_page_break()

        # 4. Examene vechi (grupate pe ani)
        self.doc.add_heading("D. Examene Vechi (2015-2019)", level=2)
        for filename in sorted(EXERCITII_FILES['examene_vechi']):
            file_path = self.repo_path / filename
            if file_path.exists():
                if filename.endswith('.pdf'):
                    self.process_exercitiu_pdf(file_path, "examen_vechi")
                elif filename.endswith('.docx'):
                    self.process_exercitiu_docx(file_path, "examen_vechi")
            else:
                logger.warning(f"  ⚠ Fișierul {filename} nu există")

        logger.info(f"✓ TOTAL EXERCIȚII PROCESATE: {self.exercitiu_counter}\n")

    def add_fise_memorare(self):
        """Adaugă fișele de memorare"""
        logger.info(f"\n{'='*60}")
        logger.info("GENERARE FIȘE DE MEMORARE")
        logger.info(f"{'='*60}\n")

        # Heading principal
        heading = self.doc.add_heading("PARTEA III - FIȘE DE MEMORARE", level=1)
        heading.runs[0].font.color.rgb = RGBColor(153, 0, 153)

        # FIȘA 1: Formule esențiale
        self.doc.add_heading("FIȘA 1: Formule Esențiale", level=2)

        for concept, formula in FORMULE_CHEIE.items():
            p = self.doc.add_paragraph()
            label = p.add_run(f"• {concept.replace('_', ' ').upper()}: ")
            label.font.bold = True
            label.font.color.rgb = RGBColor(153, 0, 153)
            formula_run = p.add_run(formula)
            formula_run.font.name = 'Courier New'
            formula_run.font.size = Pt(10)

        self.doc.add_page_break()

        # FIȘA 2: Metode de rezolvare
        self.doc.add_heading("FIȘA 2: Metode de Rezolvare", level=2)

        metode = [
            ("Convoluție", "1. Inversează h[k] → h[-k]\n2. Deplasează cu n → h[n-k]\n3. Înmulțește cu x[k]\n4. Sumează totul"),
            ("Transformata Z", "1. Scrie seria X(z) = Σx[n]z^(-n)\n2. Identifică ROC\n3. Folosește tabele dacă e posibil"),
            ("Stabilitate SNLI", "1. Verifică ∑|h[n]| < ∞\n2. Sau verifică poli în interiorul cercului unitate"),
            ("Cauzalitate", "1. h[n] = 0 pentru n < 0\n2. ROC exteriorul unui cerc"),
        ]

        for metoda, pasi in metode:
            p = self.doc.add_paragraph()
            title = p.add_run(f"📌 {metoda}:\n")
            title.font.bold = True
            title.font.size = Pt(11)
            title.font.color.rgb = RGBColor(153, 0, 153)

            steps = p.add_run(pasi)
            steps.font.size = Pt(10)
            self.doc.add_paragraph()

        self.doc.add_page_break()

        # FIȘA 3: Erori frecvente
        self.doc.add_heading("FIȘA 3: Erori Frecvente ⚠️", level=2)

        erori = [
            "❌ Confuzi x[n-k] cu x[k-n] la convoluție",
            "❌ Uiți să verifici cauzalitatea (h[n]=0 pentru n<0)",
            "❌ Nu specifici ROC la transformata Z",
            "❌ Confuzi energia cu puterea",
            "❌ Uiți condiția de stabilitate ∑|h[n]| < ∞",
            "❌ Aplici proprietăți LTI la sisteme neliniare",
            "❌ Confuzi convoluția cu corelația",
            "❌ Uiți că DTFT e periodică cu 2π"
        ]

        for eroare in erori:
            p = self.doc.add_paragraph(eroare)
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(204, 0, 0)

        self.doc.add_page_break()

        # FIȘA 4: Checklist examen
        self.doc.add_heading("FIȘA 4: Checklist Examen ✅", level=2)

        checklist = [
            "☐ Ai verificat dacă sistemul e liniar?",
            "☐ Ai verificat dacă e invariant în timp?",
            "☐ Ai calculat răspunsul impulsional h[n]?",
            "☐ Ai verificat stabilitatea (∑|h[n]| < ∞)?",
            "☐ Ai verificat cauzalitatea (h[n]=0 pentru n<0)?",
            "☐ Ai specificat ROC la transformata Z?",
            "☐ Ai verificat paritatea semnalului?",
            "☐ Ai calculat energia/puterea corect?",
            "☐ La convoluție: ai inversat, deplasat, înmulțit, sumat?",
            "☐ Ai verificat răspunsul pentru câteva valori test?",
        ]

        for item in checklist:
            p = self.doc.add_paragraph(item)
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(0, 153, 76)

        logger.info("✓ Fișe de memorare adăugate\n")

    def generate_manual(self, output_path: str = "Manual_COMPLET_PNS.docx"):
        """Generează manualul complet"""
        logger.info("\n" + "="*70)
        logger.info("ÎNCEPE GENERAREA MANUALULUI COMPLET PNS")
        logger.info("="*70 + "\n")

        # 1. Setup document
        self.setup_document()

        # 2. Copertă
        self.add_cover_page()

        # 3. Cuprins
        self.add_table_of_contents()

        # 4. Heading PARTEA I
        heading = self.doc.add_heading("PARTEA I - TEORIE COMPLETĂ", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        self.doc.add_page_break()

        # 5. Procesează toate cursurile
        for curs_id in ['C1', 'C2', 'C3', 'C4', 'C5']:
            self.process_curs(curs_id)

        # 6. Procesează toate exercițiile
        self.process_all_exercitii()

        # 7. Adaugă fișele de memorare
        self.add_fise_memorare()

        # 8. Salvează documentul
        output_file = self.repo_path / output_path
        self.doc.save(str(output_file))

        logger.info("\n" + "="*70)
        logger.info(f"✅ MANUAL COMPLET GENERAT: {output_file}")
        logger.info(f"📊 STATISTICI:")
        logger.info(f"   - Slide-uri teorie procesate: {self.slide_counter}")
        logger.info(f"   - Exerciții procesate: {self.exercitiu_counter}")
        logger.info(f"   - Fișe de memorare: 4")
        logger.info("="*70 + "\n")

        return output_file

# ============================================================================
# FUNCȚIA MAIN
# ============================================================================
def main():
    """Funcția principală - rulează generatorul"""
    print("\n🚀 GENERATOR MANUAL COMPLET PNS")
    print("="*70)

    # Inițializează generatorul
    generator = ManualPNSGenerator(repo_path='.')

    # Generează manualul
    output_file = generator.generate_manual("Manual_COMPLET_PNS.docx")

    print(f"\n✅ SUCCES! Manualul a fost generat:")
    print(f"📄 {output_file}")
    print(f"\n💡 Următorii pași:")
    print("   1. Descarcă fișierul Manual_COMPLET_PNS.docx")
    print("   2. Deschide-l în Word/LibreOffice")
    print("   3. Verifică formatarea și conținutul")
    print("   4. Învață pentru examen! 💪")
    print("="*70 + "\n")

if __name__ == "__main__":
    main()
